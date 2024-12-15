"""
Generates a comprehensive vulnerability report in both DOCX and PDF formats.
Includes sections for subdomains, vulnerabilities, open ports, SSL/TLS issues, and SQLMap results.
"""

from docx import Document
import pdfkit
from datetime import datetime
import os

def generate_report(url, subdomains, vulnerabilities, open_ports, ssl_issues, sqlmap_results):
    # Create 'reports' folder if it doesn't exist
    reports_dir = "reports"
    if not os.path.exists(reports_dir):
        os.makedirs(reports_dir)

    # Process the domain name from the URL
    domain = url.replace("https://", "").replace("http://", "").replace("/", "")

    # Add timestamp to the file name to avoid overwriting
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_docx_path = os.path.join(reports_dir, f"{domain}_vulnerability_report_{timestamp}.docx")
    report_pdf_path = os.path.join(reports_dir, f"{domain}_vulnerability_report_{timestamp}.pdf")

    # Create a new Word Document
    doc = Document()
    doc.add_heading(f"Vulnerability Report for {domain}", 0)

    # Subdomains Section
    doc.add_heading("Subdomains Identified", level=1)
    if subdomains:
        doc.add_paragraph("\n".join(subdomains))
    else:
        doc.add_paragraph("No subdomains found.")

    # Vulnerabilities Section
    doc.add_heading("Vulnerabilities Found", level=1)
    if vulnerabilities:
        for vuln in vulnerabilities:
            doc.add_paragraph(f"Type: {vuln['type']}\nDetail: {vuln['detail']}")
    else:
        doc.add_paragraph("No vulnerabilities found.")

    # Open Ports Section
    doc.add_heading("Open Ports", level=1)
    if open_ports:
        for host, ports in open_ports.items():
            port_details = ", ".join([str(port) for port in ports])
            doc.add_paragraph(f"Host: {host}\nPorts: {port_details}")
    else:
        doc.add_paragraph("No open ports found.")

    # SSL Issues Section
    doc.add_heading("SSL Issues", level=1)
    if ssl_issues:
        for issue in ssl_issues:
            doc.add_paragraph(issue)
    else:
        doc.add_paragraph("No SSL/TLS issues found.")

    # SQLMap Results Section
    doc.add_heading("SQLMap Results", level=1)
    if sqlmap_results:
        for result in sqlmap_results:
            doc.add_paragraph(result)
    else:
        doc.add_paragraph("No SQLMap results found.")

    # Save the DOCX report
    doc.save(report_docx_path)

    # Convert DOCX to PDF using pdfkit
    try:
        pdfkit.from_file(report_docx_path, report_pdf_path)
    except Exception as e:
        print(f"[!] Error generating PDF: {e}")

    return report_docx_path, report_pdf_path
